import docx
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = docx.Document()

# Setup styles
style_normal = doc.styles['Normal']
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(16)
# Right to left for Arabic
# doc.styles['Normal'].paragraph_format.bidi = True

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    return p

def add_paragraph(text, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY):
    p = doc.add_paragraph(text)
    p.alignment = align
    return p

# --- Cover Page ---
add_paragraph("وزارة التعليم العالي", WD_PARAGRAPH_ALIGNMENT.CENTER)
add_paragraph("أكاديمية الدلتا للعلوم والتكنولوجيا", WD_PARAGRAPH_ALIGNMENT.CENTER)
add_paragraph("معهد الدلتا العالي لنظم المعلومات الإدارية والمحاسبية بالمنصورة", WD_PARAGRAPH_ALIGNMENT.CENTER)

doc.add_picture('extracted_images/img_00.png', width=Inches(2.0))
doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

for _ in range(3): add_paragraph("")

title = add_paragraph("مشروع رقم (44)\nمنصة الكترونية لدعم المؤسسات استعداداً لتطبيق معايير الجودة\n(Ayn Platform)", WD_PARAGRAPH_ALIGNMENT.CENTER)
title.runs[0].font.size = Pt(24)
title.runs[0].bold = True

for _ in range(3): add_paragraph("")

sub = add_paragraph("مشروع التخرج مقدم كمتطلب جزئي للحصول علي درجة البكالوريوس في نظم معلومات الاعمال", WD_PARAGRAPH_ALIGNMENT.CENTER)
sub.runs[0].font.size = Pt(18)

for _ in range(3): add_paragraph("")

add_paragraph("تحت اشراف :", WD_PARAGRAPH_ALIGNMENT.CENTER).runs[0].bold = True
add_paragraph("د. اسماء عبد الباسط\nم. نورهان مراد", WD_PARAGRAPH_ALIGNMENT.CENTER)

add_paragraph("2026 - 2025", WD_PARAGRAPH_ALIGNMENT.CENTER)

doc.add_page_break()

# --- Quran ---
add_paragraph("بسم الله الرحمن الرحيم", WD_PARAGRAPH_ALIGNMENT.CENTER).runs[0].font.size = Pt(22)
add_paragraph("﴿ وَقُل رَّبِّ زِدْنِي عِلْمًا ﴾", WD_PARAGRAPH_ALIGNMENT.CENTER).runs[0].font.size = Pt(24)
add_paragraph("صدق الله العظيم\n(طه: 114)", WD_PARAGRAPH_ALIGNMENT.CENTER)
doc.add_page_break()

# --- Dedication ---
add_heading("إهداء", 1)
add_paragraph("إلى من ضحوا من أجلنا... إلى عائلاتنا وأساتذتنا الكرام.")
doc.add_page_break()

# --- Acknowledgements ---
add_heading("شكر وتقدير", 1)
add_paragraph("نتقدم بخالص الشكر والتقدير إلى د. محمد ربيع ناصر رئيس مجلس إدارة الأكاديمية، وإلى أ.د. أحمد أبو الفتوح عميد المعهد. كما نتوجه بالشكر الخاص إلى مشرفي المشروع د. أسماء عبد الباسط و م. نورهان مراد على توجيهاتهم ودعمهم المستمر.")
doc.add_page_break()

# --- Chapter 1: Introduction ---
add_heading("الفصل الأول: الإطار العام للمشروع (Introduction)", 1)
add_heading("1.1 مقدمة", 2)
add_paragraph("منصة Ayn هي حل مبتكر مصمم لتمكين المؤسسات التعليمية من تحقيق والحفاظ على معايير الجودة، وضمان الجاهزية للاعتماد، ودفع التحسين التشغيلي المستمر. تعمل Ayn كنظام إدارة شامل يبسط العمليات المعقدة المرتبطة بضمان الجودة الأكاديمية والامتثال. يوفر بيئة موحدة مدعومة بالذكاء الاصطناعي حيث يمكن للمؤسسات إدارة أطر الاعتماد، ومستودعات الأدلة، وسير عمل الجودة بكفاءة وذكاء غير مسبوقين.")

add_heading("1.2 تعريف المشكلة", 2)
add_paragraph("غالباً ما تتسم عمليات الاعتماد الأكاديمي وضمان الجودة التقليدية بمهام يدوية كثيفة العمالة، وبيانات مجزأة، ونقص في الرؤى في الوقت الفعلي. تكافح المؤسسات في جمع وتنظيم وربط كميات هائلة من الأدلة، مثل الوثائق والسياسات والسجلات، بمعايير اعتماد محددة. هذه العملية اليدوية لجمع الأدلة وربطها تستغرق وقتاً طويلاً وعرضة للأخطاء. وتتصدى Ayn لهذه المشاكل من خلال مركزية جميع أنشطة ضمان الجودة، وأتمتة العمليات الرئيسية، والاستفادة من الذكاء الاصطناعي لتقديم رؤى وتوصيات ذكية.")

add_heading("1.3 الأهداف", 2)
add_paragraph("1. أتمتة إدارة الأدلة وتقليل الجهد البشري.\n2. توفير تحليل الفجوات مدعوم بالذكاء الاصطناعي.\n3. تقديم رؤى في الوقت الفعلي حول حالة الامتثال.\n4. تيسير التعاون عبر الأقسام لضمان الجودة.")

doc.add_page_break()

# Read the rest of the text from the previous document
import re
with open('/tmp/docx_full.txt', 'r', encoding='utf-8') as f:
    lines = f.readlines()

current_img_idx = 1
for line in lines:
    line = line.strip()
    if not line: continue
    
    match = re.match(r'\[\d+\|([^\]]+)\] (.*)', line)
    if match:
        style = match.group(1)
        text = match.group(2)
        
        if style.startswith('Heading 1'):
            doc.add_page_break()
            add_heading(text, 1)
        elif style.startswith('Heading 2'):
            add_heading(text, 2)
        elif style.startswith('Heading 3'):
            add_heading(text, 3)
        elif style.startswith('Heading 4'):
            add_heading(text, 4)
        elif style == 'Title':
            pass # ignore
        else:
            add_paragraph(text)
            
            # Inject images periodically
            if 'صورة' in text or 'diagram' in text.lower() or 'architecture' in text.lower() or len(text) > 300:
                img_path = f'extracted_images/img_{current_img_idx:02d}.png'
                if os.path.exists(img_path) and current_img_idx < 37:
                    try:
                        doc.add_picture(img_path, width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        current_img_idx += 1
                    except Exception as e:
                        print(f"Failed to add image {img_path}: {e}")

doc.save('Ayn_Final_Book.docx')
print("Document generated successfully as Ayn_Final_Book.docx")
