"""
Build AynPlatformGraduationBook_Formatted.docx
Content: AynPlatformGraduationBook.docx
Format: المشروع تعديلات 28-4-2026.pdf
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import zipfile, os, shutil, io

SRC_DOCX = '/home/anwar/Projects/Ayn/grad/AynPlatformGraduationBook.docx'
OUT_DOCX = '/home/anwar/Projects/Ayn/grad/AynPlatformGraduationBook_Formatted.docx'
IMG_DIR  = '/home/anwar/Projects/Ayn/grad/docx_images'

# ── helpers ──────────────────────────────────────────────────────────────────

def set_rtl(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)

def add_run_fmt(para, text, font_name, size_pt, bold=False, color=None, rtl=False):
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # set rFonts for Arabic
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)
    if rtl:
        rtl_el = OxmlElement('w:rtl')
        rPr.append(rtl_el)
    return run

def set_para_spacing(para, before=0, after=6, line=None):
    pPr = para._p.get_or_add_pPr()
    pPr_spacing = OxmlElement('w:spacing')
    pPr_spacing.set(qn('w:before'), str(int(before * 20)))
    pPr_spacing.set(qn('w:after'), str(int(after * 20)))
    if line:
        pPr_spacing.set(qn('w:line'), str(int(line * 240)))
        pPr_spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(pPr_spacing)

def set_page_setup(doc):
    section = doc.sections[0]
    # A4
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    # Margins matching reference (approx 2.5cm each side)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type())
    return para

def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br  # Not used directly

def insert_page_break(doc):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    doc.element.body.append(p)

def add_heading_para(doc, text, level, arabic=True):
    """Add a heading with reference-PDF style"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.RIGHT

    if level == 1:
        # Chapter title - SimplifiedArabic-Bold 24pt centered
        add_run_fmt(para, text, 'Simplified Arabic', 20, bold=True, rtl=arabic)
        set_para_spacing(para, before=12, after=12)
    elif level == 2:
        # Section heading - SimplifiedArabic-Bold 18pt
        add_run_fmt(para, text, 'Simplified Arabic', 16, bold=True, rtl=arabic)
        set_para_spacing(para, before=10, after=6)
    elif level == 3:
        add_run_fmt(para, text, 'Simplified Arabic', 14, bold=True, rtl=arabic)
        set_para_spacing(para, before=8, after=4)
    else:
        add_run_fmt(para, text, 'Simplified Arabic', 13, bold=True, rtl=arabic)
        set_para_spacing(para, before=6, after=3)

    if arabic:
        set_rtl(para)
    return para

def add_body_para(doc, text, arabic=True):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if not arabic else WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run_fmt(para, text, 'Simplified Arabic' if arabic else 'Calibri', 14, rtl=arabic)
    set_para_spacing(para, before=0, after=6, line=1.5)
    if arabic:
        set_rtl(para)
    return para

def add_image_centered(doc, img_path, max_width_cm=15):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    try:
        run.add_picture(img_path, width=Cm(max_width_cm))
    except Exception as e:
        para.add_run(f'[Image: {os.path.basename(img_path)}]')
    set_para_spacing(para, before=6, after=6)
    return para

def add_cover_page(doc):
    """Cover page matching reference PDF page 1 style"""
    # University name
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fmt(p1, 'وزارة التعليم العالي', 'Simplified Arabic', 16, bold=True, rtl=True)
    set_rtl(p1)
    set_para_spacing(p1, before=60, after=4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fmt(p2, 'أكاديمية الدلتا للعلوم والتكنولوجيا', 'Simplified Arabic', 16, bold=True, rtl=True)
    set_rtl(p2)
    set_para_spacing(p2, before=4, after=4)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fmt(p3, 'معهد الدلتا العالي لنظم المعلومات الإدارية والمحاسبية', 'Simplified Arabic', 14, bold=True, rtl=True)
    set_rtl(p3)
    set_para_spacing(p3, before=4, after=30)

    # Logo image (image1.png from DOCX)
    logo_path = os.path.join(IMG_DIR, 'image1.png')
    if os.path.exists(logo_path):
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.add_run().add_picture(logo_path, width=Cm(5))
        set_para_spacing(lp, before=10, after=10)

    # Platform title
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fmt(pt, 'منصة عين (Ayn)', 'Simplified Arabic', 22, bold=True, rtl=True)
    set_rtl(pt)
    set_para_spacing(pt, before=20, after=6)

    pt2 = doc.add_paragraph()
    pt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fmt(pt2, 'منصة إلكترونية لدعم المؤسسات استعداداً لتطبيق معايير الجودة', 'Simplified Arabic', 16, bold=False, rtl=True)
    set_rtl(pt2)
    set_para_spacing(pt2, before=4, after=30)

    # Year
    py = doc.add_paragraph()
    py.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fmt(py, '2025 - 2026', 'Calibri', 14)
    set_para_spacing(py, before=20, after=4)

def add_quran_page(doc):
    insert_page_break(doc)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fmt(p, 'بِسْمِ اللَّهِ الرَّحْمَنِ الرَّحِيمِ', 'Simplified Arabic', 18, bold=True, rtl=True)
    set_rtl(p)
    set_para_spacing(p, before=20, after=20)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fmt(p2, 'وَقُل رَّبِّ زِدْنِي عِلْمًا', 'Simplified Arabic', 18, rtl=True)
    set_rtl(p2)
    set_para_spacing(p2, before=10, after=10)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fmt(p3, 'صَدَقَ اللَّهُ الْعَظِيمُ', 'Simplified Arabic', 16, rtl=True)
    set_rtl(p3)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fmt(p4, '(طه: 114)', 'Simplified Arabic', 14, rtl=True)
    set_rtl(p4)

def add_dedication_page(doc):
    insert_page_break(doc)
    add_heading_para(doc, 'الإهداء', 1)

def add_acknowledgment_page(doc):
    insert_page_break(doc)
    add_heading_para(doc, 'شكر وتقدير', 1)

def add_abstract_pages(doc, src_doc):
    """Add abstract content from source doc"""
    insert_page_break(doc)
    add_heading_para(doc, 'الملخص', 1)

def add_toc_page(doc):
    insert_page_break(doc)
    add_heading_para(doc, 'فهرس المحتويات', 1)

def add_chapter_divider(doc, chapter_num_ar, chapter_title_ar, chapter_title_en):
    insert_page_break(doc)
    # Blank space then chapter number/title centered - matches reference PDF style
    for _ in range(8):
        doc.add_paragraph()
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fmt(p1, f'{chapter_num_ar}', 'Simplified Arabic', 24, bold=True, rtl=True)
    set_rtl(p1)
    set_para_spacing(p1, before=10, after=6)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fmt(p2, chapter_title_ar, 'Simplified Arabic', 24, bold=True, rtl=True)
    if chapter_title_en:
        add_run_fmt(p2, f'  ({chapter_title_en})', 'Times New Roman', 20, bold=True)
    set_rtl(p2)
    set_para_spacing(p2, before=6, after=10)

def is_arabic(text):
    for c in text:
        if '\u0600' <= c <= '\u06FF':
            return True
    return False

# ── main build ────────────────────────────────────────────────────────────────

def build():
    src = Document(SRC_DOCX)
    out = Document()
    set_page_setup(out)

    # Remove default empty paragraph
    for p in out.paragraphs:
        p._element.getparent().remove(p._element)

    # -- Load image map from src doc
    rels = src.part.rels
    rId_to_img = {}
    for rId, rel in rels.items():
        if 'image' in rel.reltype.lower():
            fname = str(rel.target_ref).split('/')[-1]
            rId_to_img[rId] = fname

    def get_para_image(para):
        for run in para.runs:
            for pic in run.element.findall('.//' + qn('a:blip')):
                rId = pic.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', '')
                fname = rId_to_img.get(rId, '')
                full = os.path.join(IMG_DIR, fname)
                if os.path.exists(full):
                    return full
        return None

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    add_cover_page(out)

    # Cover image (image20.png is the large cover image in source doc para 2)
    cover_img = os.path.join(IMG_DIR, 'image20.png')
    if os.path.exists(cover_img):
        p = out.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(cover_img, width=Cm(14))
        set_para_spacing(p, before=10, after=10)

    # ── QURAN VERSE ───────────────────────────────────────────────────────────
    add_quran_page(out)

    # ── DEDICATION ────────────────────────────────────────────────────────────
    add_dedication_page(out)
    # Extract dedication text from source (paragraphs 3-25 area before heading)
    for i, para in enumerate(src.paragraphs[3:26]):
        t = para.text.strip()
        if t:
            add_body_para(out, t, arabic=is_arabic(t))

    # ── ACKNOWLEDGMENT ────────────────────────────────────────────────────────
    add_acknowledgment_page(out)
    for i, para in enumerate(src.paragraphs[26:60]):
        t = para.text.strip()
        style = para.style.name
        if not t:
            continue
        img = get_para_image(para)
        if img:
            add_image_centered(out, img, max_width_cm=6)
        elif 'Heading' in style:
            add_heading_para(out, t, 2, arabic=is_arabic(t))
        else:
            add_body_para(out, t, arabic=is_arabic(t))

    # ── ABSTRACT (Arabic) ─────────────────────────────────────────────────────
    insert_page_break(out)
    add_heading_para(out, 'الملخص', 1)
    for para in src.paragraphs[60:100]:
        t = para.text.strip()
        style = para.style.name
        if not t:
            continue
        img = get_para_image(para)
        if img:
            add_image_centered(out, img)
        elif 'Heading' in style:
            add_heading_para(out, t, 2, arabic=is_arabic(t))
        else:
            add_body_para(out, t, arabic=is_arabic(t))

    # ── TABLE OF CONTENTS (placeholder) ───────────────────────────────────────
    add_toc_page(out)
    toc_entries = [
        ('الفصل الأول: الإطار العام للمشروع (Introduction)', '1'),
        ('الفصل الثاني: فكرة المشروع (Project Idea)', '2'),
        ('الفصل الثالث: الأدوات والتقنيات (Tools & Technologies)', '3'),
        ('الفصل الرابع: تحليل النظام (System Analysis & Architecture)', '4'),
        ('الفصل الخامس: التنفيذ (Implementation)', '5'),
        ('الفصل السادس: تصميم الواجهات (UI/UX Design)', '6'),
        ('الفصل السابع: الاختبار والنتائج (Testing & Results)', '7'),
        ('المراجع والملاحق (References & Appendices)', '8'),
    ]
    for entry, pg in toc_entries:
        p = out.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run_fmt(p, entry, 'Simplified Arabic', 14, rtl=True)
        add_run_fmt(p, f'  ............  {pg}', 'Calibri', 12)
        set_rtl(p)
        set_para_spacing(p, before=4, after=4)

    # ── CHAPTER PROCESSING ────────────────────────────────────────────────────
    # Map source heading structure to chapter ranges
    chapters = [
        {
            'num': 'الفصل الأول',
            'title_ar': 'الإطار العام للمشروع',
            'title_en': 'Introduction',
            'start': 159,
            'end': 265,
        },
        {
            'num': 'الفصل الثاني',
            'title_ar': 'فكرة المشروع',
            'title_en': 'Project Idea',
            'start': 266,
            'end': 361,
        },
        {
            'num': 'الفصل الثالث',
            'title_ar': 'الأدوات والتقنيات',
            'title_en': 'Tools & Technologies',
            'start': 362,
            'end': 399,
        },
        {
            'num': 'الفصل الرابع',
            'title_ar': 'تحليل النظام والمعمارية',
            'title_en': 'System Analysis & Architecture',
            'start': 400,
            'end': 541,
        },
        {
            'num': 'الفصل الخامس',
            'title_ar': 'التنفيذ',
            'title_en': 'Implementation',
            'start': 542,
            'end': 652,
        },
        {
            'num': 'الفصل السادس',
            'title_ar': 'تصميم واجهات المستخدم',
            'title_en': 'UI/UX Design',
            'start': 653,
            'end': 767,
        },
        {
            'num': 'الفصل السابع',
            'title_ar': 'الاختبار والنتائج',
            'title_en': 'Testing & Results',
            'start': 768,
            'end': 801,
        },
        {
            'num': 'المراجع والملاحق',
            'title_ar': 'المراجع والملاحق',
            'title_en': 'References & Appendices',
            'start': 802,
            'end': 824,
        },
    ]

    for chap in chapters:
        add_chapter_divider(out, chap['num'], chap['title_ar'], chap['title_en'])
        insert_page_break(out)

        paras = src.paragraphs[chap['start']:chap['end']]
        for para in paras:
            t = para.text.strip()
            style = para.style.name
            img = get_para_image(para)

            if img:
                add_image_centered(out, img, max_width_cm=14)
                if t:
                    # caption
                    cap = out.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_run_fmt(cap, t, 'Simplified Arabic', 11, rtl=is_arabic(t))
                    set_para_spacing(cap, before=2, after=8)
                continue

            if not t:
                continue

            arabic = is_arabic(t)

            if 'Heading 1' in style or style == 'Title':
                add_heading_para(out, t, 1, arabic=arabic)
            elif 'Heading 2' in style:
                add_heading_para(out, t, 2, arabic=arabic)
            elif 'Heading 3' in style:
                add_heading_para(out, t, 3, arabic=arabic)
            elif 'Heading 4' in style:
                add_heading_para(out, t, 4, arabic=arabic)
            else:
                add_body_para(out, t, arabic=arabic)

    out.save(OUT_DOCX)
    print(f"✅ Saved: {OUT_DOCX}")

if __name__ == '__main__':
    build()
